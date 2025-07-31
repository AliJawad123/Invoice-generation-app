import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

st.set_page_config(page_title="Quotation Generator", layout="centered")
st.title("Quotation Generator App")

# --- Quotation Details ---
st.header("Quotation Details")
invoice_number = st.text_input("Quotation Number")
invoice_date = st.text_input("Quotation Date")
company_name = st.text_input("Company Name")
company_logo = st.file_uploader("Upload Company Logo (optional)", type=["png", "jpg", "jpeg"])

# --- Quotation From ---
st.header("Quotation From")
invoice_from_name = st.text_input("Name (Quotation From)")
invoice_from_email = st.text_input("Email (Quotation From)")
invoice_from_contact = st.text_input("Contact Number (Quotation From)")

# --- Quotation To ---
st.header("Quotation To")
invoice_to_name = st.text_input("Name (Quotation To)")
invoice_to_email = st.text_input("Email (Quotation To)")
invoice_to_contact = st.text_input("Contact Number (Quotation To)")

# --- Items ---
st.header("Items")
items = []
num_items = st.number_input("Number of Items", min_value=1, value=1)

for i in range(int(num_items)):
    st.subheader(f"Item {i+1}")
    serial_number = st.text_input(f"Serial Number {i+1}", key=f"serial_{i}")
    production_description = st.text_input(f"Production Description {i+1}", key=f"desc_{i}")
    quantity = st.number_input(f"Quantity {i+1}", min_value=1, value=1, key=f"qty_{i}")
    price = st.number_input(f"Price {i+1} (PKR)", min_value=0.0, value=0.0, key=f"price_{i}")
    total = quantity * price
    items.append({
        "serial_number": serial_number,
        "production_description": production_description,
        "quantity": quantity,
        "price": price,
        "total": total
    })

# --- Additional Charges ---
st.header("Additional Charges")
shipping_charges = st.number_input("Shipping Charges (PKR)", min_value=0.0, value=0.0)
packaging_charges = st.number_input("Packaging Charges (PKR)", min_value=0.0, value=0.0)
tax_rate = st.number_input("Tax Rate (%)", min_value=0.0, value=0.0)


# --- Validation ---
def check_empty_fields():
    empty_fields = []
    if not invoice_number: empty_fields.append("Quotation Number")
    if not invoice_date: empty_fields.append("Quotation Date")
    if not company_name: empty_fields.append("Company Name")
    if not invoice_from_name: empty_fields.append("Name (Quotation From)")
    if not invoice_from_email: empty_fields.append("Email (Quotation From)")
    if not invoice_from_contact: empty_fields.append("Contact Number (Quotation From)")
    if not invoice_to_name: empty_fields.append("Name (Quotation To)")
    if not invoice_to_email: empty_fields.append("Email (Quotation To)")
    if not invoice_to_contact: empty_fields.append("Contact Number (Quotation To)")
    for i, item in enumerate(items):
        if not item["serial_number"]: empty_fields.append(f"Serial Number {i+1}")
        if not item["production_description"]: empty_fields.append(f"Production Description {i+1}")
        if item["quantity"] <= 0: empty_fields.append(f"Quantity {i+1}")
        if item["price"] <= 0: empty_fields.append(f"Price {i+1}")
    return empty_fields


# --- Word Document Generator ---
def generate_invoice_docx(invoice_number, invoice_date, company_name, company_logo,
                          invoice_from_name, invoice_from_email, invoice_from_contact,
                          invoice_to_name, invoice_to_email, invoice_to_contact,
                          items, shipping_charges, packaging_charges, tax_rate):

    doc = Document()

    # Company Name at the top
    header = doc.add_paragraph()
    run = header.add_run(company_name.upper())
    run.bold = True
    run.font.size = Pt(24)
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add logo if available
    if company_logo:
        doc.add_picture(company_logo, width=Inches(1.5))
    
    doc.add_paragraph(f"Quotation Number: {invoice_number}")
    doc.add_paragraph(f"Quotation Date: {invoice_date}")

    doc.add_paragraph("")

    doc.add_paragraph("Quotation From:")
    doc.add_paragraph(f"{invoice_from_name}")
    doc.add_paragraph(f"Email: {invoice_from_email}")
    doc.add_paragraph(f"Contact: {invoice_from_contact}")

    doc.add_paragraph("")

    doc.add_paragraph("Quotation To:")
    doc.add_paragraph(f"{invoice_to_name}")
    doc.add_paragraph(f"Email: {invoice_to_email}")
    doc.add_paragraph(f"Contact: {invoice_to_contact}")

    doc.add_paragraph("")

    # Add Table for Items
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'S.No.'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Qty'
    hdr_cells[3].text = 'Price (PKR)'
    hdr_cells[4].text = 'Total (PKR)'

    total_amount = 0
    for item in items:
        row_cells = table.add_row().cells
        row_cells[0].text = item["serial_number"]
        row_cells[1].text = item["production_description"]
        row_cells[2].text = str(item["quantity"])
        row_cells[3].text = f"{item['price']:.2f}"
        row_cells[4].text = f"{item['total']:.2f}"
        total_amount += item["total"]

    tax = total_amount * (tax_rate / 100)
    total_incl_tax = total_amount + shipping_charges + packaging_charges + tax

    doc.add_paragraph("")
    doc.add_paragraph(f"Subtotal: {total_amount:.2f} PKR")
    doc.add_paragraph(f"Shipping Charges: {shipping_charges:.2f} PKR")
    doc.add_paragraph(f"Packaging Charges: {packaging_charges:.2f} PKR")
    doc.add_paragraph(f"Tax ({tax_rate}%): {tax:.2f} PKR")
    doc.add_paragraph(f"Total (Incl. Tax): {total_incl_tax:.2f} PKR")

    doc.add_paragraph("")
    doc.add_paragraph("All prices are in Pakistani Rupees (PKR).")

    # Return Word file as bytes
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# --- Generate Word Button ---
if st.button("Generate Quotation Word File"):
    empty_fields = check_empty_fields()
    if not empty_fields:
        docx_bytes = generate_invoice_docx(
            invoice_number, invoice_date, company_name, company_logo,
            invoice_from_name, invoice_from_email, invoice_from_contact,
            invoice_to_name, invoice_to_email, invoice_to_contact,
            items, shipping_charges, packaging_charges, tax_rate
        )
        st.success("Quotation Word file generated successfully!")
        st.download_button(
            label="📄 Download Word File",
            data=docx_bytes,
            file_name="Quotation.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.error(f"Please fill out the following fields: {', '.join(empty_fields)}")
